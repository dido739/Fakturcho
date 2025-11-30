import re
import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_amount_token(token):
    if not token:
        return None
    t = token.strip()
    neg = False
    if t.startswith('(') and t.endswith(')'):
        neg = True
        t = t[1:-1].strip()
    if t.startswith('-'):
        neg = True
        t = t[1:].strip()
    t = t.replace('\u00A0', '')
    t = t.replace('\xa0', '')
    t = t.replace(' ', '')
    if ',' in t and '.' in t:
        t = t.replace(',', '')
    elif ',' in t and '.' not in t:
        t = t.replace(',', '.')
    m = re.match(r"^[-+]?[0-9]*\.?[0-9]+", t)
    if not m:
        return None
    try:
        val = float(m.group(0))
    except ValueError:
        return None
    return -val if neg else val


def extract_euro_amounts(text):
    pattern = re.compile(r'([\d\s\u00A0\xa0\.,()\-+]+)\s*(€|\u20AC)')
    amounts = []
    for m in pattern.finditer(text):
        raw = m.group(1).strip()
        token = raw.split()[-1] if raw.split() else raw
        val = parse_amount_token(token)
        if val is not None:
            amounts.append(val)
    return amounts


def extract_iban(text):
    pattern = re.compile(r'IBAN[:\s]*([A-Z]{2}[0-9A-Z]+)', re.IGNORECASE)
    m = pattern.search(text)
    if m:
        return m.group(1).strip()
    pattern2 = re.compile(r'\b(BG[0-9]{2}[A-Z0-9]+)\b')
    m2 = pattern2.search(text)
    if m2:
        return m2.group(1).strip()
    return None


def extract_poluchatel(text):
    pattern = re.compile(r'Получател[:\s]+([^\n]+)', re.IGNORECASE)
    m = pattern.search(text)
    if m:
        return m.group(1).strip()
    return None


def extract_dostavchik(text):
    pattern = re.compile(r'Доставчик[:\s]+([^\n]+)', re.IGNORECASE)
    m = pattern.search(text)
    if m:
        return m.group(1).strip()
    return None


def extract_field_after_label(text, label):
    pattern = re.compile(rf'{label}[:\s№\.]*([^\n]+)', re.IGNORECASE)
    m = pattern.search(text)
    if m:
        value = m.group(1).strip()
        value = re.split(r'\b(ДДС|Идент|Град|Адрес|МОЛ|Дата|Номер|Получател|Доставчик)\b', value, flags=re.IGNORECASE)[0].strip()
        return value if value else None
    return None


def extract_description(text):
    pattern = re.search(r'Наименование на стоката/услугата\s*([\s\S]*?)(?=\s*Сума|IBAN|$)', text, re.IGNORECASE)
    if pattern:
        desc = pattern.group(1).strip()
        return desc if desc else None
    return None


def extract_all_fields(text):
    fields = {
        'poluchatel': extract_poluchatel(text),
        'dostavchik': extract_dostavchik(text),
        'iban': extract_iban(text),
        'description': extract_description(text),
    }
    
    poluchatel_section = re.search(r'Получател(.*?)Доставчик', text, re.IGNORECASE | re.DOTALL)
    if poluchatel_section:
        section_text = poluchatel_section.group(1)
        fields['poluchatel_dds'] = extract_field_after_label(section_text, r'ДДС\s*№')
        fields['poluchatel_ident'] = extract_field_after_label(section_text, r'Идент\.?\s*№')
        fields['poluchatel_grad'] = extract_field_after_label(section_text, r'Град')
        fields['poluchatel_adres'] = extract_field_after_label(section_text, r'Адрес')
        fields['poluchatel_mol'] = extract_field_after_label(section_text, r'МОЛ')
    
    dostavchik_section = re.search(r'Доставчик(.*?)(?=Наименование|IBAN|$)', text, re.IGNORECASE | re.DOTALL)
    if dostavchik_section:
        section_text = dostavchik_section.group(1)
        fields['dostavchik_dds'] = extract_field_after_label(section_text, r'ДДС\s*№')
        fields['dostavchik_ident'] = extract_field_after_label(section_text, r'Идент\.?\s*№')
        fields['dostavchik_grad'] = extract_field_after_label(section_text, r'Град')
        fields['dostavchik_adres'] = extract_field_after_label(section_text, r'Адрес')
        fields['dostavchik_mol'] = extract_field_after_label(section_text, r'МОЛ')
    
    fields['data'] = extract_field_after_label(text, r'Дата')
    fields['nomer'] = extract_field_after_label(text, r'Номер')
    
    return fields


def format_bulgarian(num):
    s = f"{num:,.2f}"
    s = s.replace(',', 'TEMP').replace('.', ',').replace('TEMP', ' ')
    return s


def extract_euro_amounts_from_doc(doc):
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    full_text = '\n'.join(texts)
    return extract_euro_amounts(full_text)


def fill_template_table(doc, amounts, fields):
    if not amounts:
        return doc
    
    base = amounts[0]
    dds = round(base * 0.20, 2)
    total = round(base * 1.20, 2)
    
    for table in doc.tables:
        if len(table.rows) == 4 and len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.lower().strip()
            if 'наименование' in first_cell:
                table.rows[0].cells[0].text = 'Наименование на стоката/услугата'
                if fields.get('description'):
                    table.rows[0].cells[1].text = fields['description']
                else:
                    table.rows[0].cells[1].text = ''
                table.rows[1].cells[1].text = format_bulgarian(base) + " €"
                table.rows[2].cells[1].text = format_bulgarian(dds) + " €"
                table.rows[3].cells[1].text = format_bulgarian(total) + " €"
        
        if len(table.rows) == 3 and len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.lower().strip()
            if 'данъчна основа' in first_cell:
                table.rows[0].cells[1].text = format_bulgarian(base) + " €"
                table.rows[1].cells[1].text = format_bulgarian(dds) + " €"
                table.rows[2].cells[1].text = format_bulgarian(total) + " €"
        
        if fields.get('iban') and len(table.rows) == 1 and len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.lower().strip()
            if 'iban' in first_cell:
                table.rows[0].cells[1].text = fields['iban']
        
        if len(table.rows) == 6 and len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.strip()
            if first_cell == 'Получател':
                if fields.get('poluchatel'):
                    table.rows[0].cells[0].text = f"Получател: {fields['poluchatel']}"
                if fields.get('poluchatel_dds'):
                    table.rows[1].cells[0].text = f"ДДС№: {fields['poluchatel_dds']}"
                if fields.get('poluchatel_ident'):
                    table.rows[2].cells[0].text = f"Индент№: {fields['poluchatel_ident']}"
                if fields.get('poluchatel_grad'):
                    table.rows[3].cells[0].text = f"Град: {fields['poluchatel_grad']}"
                if fields.get('poluchatel_adres'):
                    table.rows[4].cells[0].text = f"Адрес: {fields['poluchatel_adres']}"
                if fields.get('poluchatel_mol'):
                    table.rows[5].cells[0].text = f"Мол: {fields['poluchatel_mol']}"
                
                if fields.get('dostavchik'):
                    table.rows[0].cells[1].text = f"Доставчик: {fields['dostavchik']}"
                if fields.get('dostavchik_dds'):
                    table.rows[1].cells[1].text = f"ДДС№: {fields['dostavchik_dds']}"
                if fields.get('dostavchik_ident'):
                    table.rows[2].cells[1].text = f"Индент№: {fields['dostavchik_ident']}"
                if fields.get('dostavchik_grad'):
                    table.rows[3].cells[1].text = f"Град: {fields['dostavchik_grad']}"
                if fields.get('dostavchik_adres'):
                    table.rows[4].cells[1].text = f"Адрес: {fields['dostavchik_adres']}"
                if fields.get('dostavchik_mol'):
                    table.rows[5].cells[1].text = f"Мол: {fields['dostavchik_mol']}"
    
    for paragraph in doc.paragraphs:
        if 'Фактура' in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                if 'Фактура' in run.text:
                    run.font.size = Pt(48)
                    run.font.bold = True
        if 'Дата:' in paragraph.text and fields.get('data'):
            paragraph.text = paragraph.text.replace('Дата:', f"Дата: {fields['data']}")
        if 'Номер:' in paragraph.text and fields.get('nomer'):
            paragraph.text = paragraph.text.replace('Номер:', f"Номер: {fields['nomer']}")
    
    return doc


def process_docx_template(template_path, output_path):
    doc = Document(template_path)
    amounts = extract_euro_amounts_from_doc(doc)
    
    if not amounts:
        print("No euro amounts found in template.")
        doc.save(output_path)
        return output_path
    
    doc = fill_template_table(doc, amounts)
    doc.save(output_path)
    return output_path


def ensure_docx_output_path(in_path, out_path=None):
    if out_path:
        return out_path
    folder = os.path.dirname(os.path.abspath(in_path)) or os.getcwd()
    base = 'calc_output'
    idx = 1
    while True:
        cand = os.path.join(folder, f"{base}{idx}.docx")
        if not os.path.exists(cand):
            return cand
        idx += 1


def main():
    if len(sys.argv) < 3:
        print("Usage: python calc_combined.py test.txt template.docx [output.docx]")
        raise SystemExit(1)
    
    txt_path = sys.argv[1]
    template_path = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not os.path.exists(txt_path):
        print(f"Text file not found: {txt_path}")
        raise SystemExit(1)
    
    if not os.path.exists(template_path):
        print(f"Template file not found: {template_path}")
        raise SystemExit(1)
    
    with open(txt_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()
    
    amounts = extract_euro_amounts(text)
    fields = extract_all_fields(text)
    
    if not amounts:
        print("No euro amounts found in text file.")
        raise SystemExit(0)
    
    doc = Document(template_path)
    doc = fill_template_table(doc, amounts, fields)
    
    
    out_path = ensure_docx_output_path(template_path, output_path)
    doc.save(out_path)
    print(f"Created DOCX file: {out_path}")


if __name__ == '__main__':
    main()
